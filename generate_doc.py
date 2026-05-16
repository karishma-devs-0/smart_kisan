import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Colors
DARK_GREEN = RGBColor(0x1B, 0x5E, 0x20) # #1B5E20
PRIMARY_GREEN = RGBColor(0x2E, 0x7D, 0x32) # #2E7D32
LIGHT_GREEN = RGBColor(0x4C, 0xAF, 0x50) # #4CAF50
ORANGE = RGBColor(0xF5, 0x7F, 0x17) # #F57F17

doc = Document()

# Add Title
title = doc.add_heading('SmartKisan Session Summary', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = DARK_GREEN
    run.font.bold = True

doc.add_paragraph()

# Custom Styles
def add_section_header(doc, text):
    p = doc.add_heading(level=1)
    run = p.add_run(text)
    run.font.color.rgb = PRIMARY_GREEN
    run.font.name = 'Arial'
    return p

def add_bullet(doc, text, bold_start=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_start:
        p.add_run(bold_start).bold = True
        p.add_run(text[len(bold_start):])
    else:
        p.add_run(text)
    return p

# Section 1
add_section_header(doc, "1. DB Schema Updates")
p = doc.add_paragraph("Moved away from Firebase to a standard relational schema:")
add_bullet(doc, "PolyCab Pattern: Switched to Postgres using auto-increment integer IDs and added deleted_at for soft deletes.", "PolyCab Pattern: ")
add_bullet(doc, "Auth: Added tables for users, otps, and refresh_tokens to handle logins ourselves.", "Auth: ")
add_bullet(doc, "Enums & Data Types: Added proper enums for roles, statuses, and device types to keep data clean.", "Enums & Data Types: ")
add_bullet(doc, "Missing Tables: Added the missing tables for settings, onboarding_profile, uploads, etc.", "Missing Tables: ")

# Section 2
add_section_header(doc, "2. Backend Merge (Rohit's Code)")
p = doc.add_paragraph("Pulled in Rohit's Express backend code to replace Firebase.")
add_bullet(doc, "Database Connection: Added schema.sql and the db connection pool.", "Database Connection: ")
add_bullet(doc, "API Endpoints: Merged the routes for auth, pumps, and pump groups.", "API Endpoints: ")
add_bullet(doc, "Packages: Installed needed packages like bcrypt, jsonwebtoken, pg, etc.", "Packages: ")
add_bullet(doc, "Env Vars: Merged the .env files for Postgres and MQTT settings.", "Env Vars: ")

# Section 3
add_section_header(doc, "3. Removing Firebase")
p = doc.add_paragraph("Stripped Firebase completely out of the app.")
add_bullet(doc, "Disabled Firebase: Flipped FIREBASE_ENABLED to false to disable SDK calls.", "Disabled Firebase: ")
add_bullet(doc, "App.js Cleanup: Removed Firebase imports and the old onAuthStateChanged listener.", "App.js Cleanup: ")
add_bullet(doc, "Auth Flow: Updated secureAuth.js and authSlice.js to just use local JWTs in secure storage.", "Auth Flow: ")

# Section 4
add_section_header(doc, "4. UI Updates & Navigation")
p = doc.add_paragraph("Blocked off screens that still need real backend data.")
add_bullet(doc, "Coming Soon Screen: Built a branded ComingSoonScreen.js to show for unfinished features.", "Coming Soon Screen: ")
add_bullet(doc, "Redirects: Updated HomeStack and SettingsStack to route 21 incomplete screens to the Coming Soon page.", "Redirects: ")
add_bullet(doc, "Working Features: Left the AI tools, core settings, and device management active.", "Working Features: ")

# Section 5
add_section_header(doc, "5. Getting it Running locally")
p = doc.add_paragraph("Here's what needs to be done next to run everything:")
add_bullet(doc, "Install Postgres 17 and set the password to '7090'.", "Install Postgres 17 ")
add_bullet(doc, "Create a database named 'smartkisan'.", "Create a database ")
add_bullet(doc, "Run 'psql -U postgres -d smartkisan -f backend/src/database/schema.sql' to create the tables.", "Run schema.sql ")
add_bullet(doc, "Run 'npm run dev' in the backend folder to start the API.", "Start the API ")
add_bullet(doc, "Update api.js in the frontend to point to the new localhost:5000 endpoints.", "Update api.js ")

# Save document
output_path = r'C:\Users\karishma\smart_kisan\SmartKisan_Session_Summary_v2.docx'
doc.save(output_path)
print(f"Document saved to {output_path}")
